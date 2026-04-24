import os
import re
from collections import defaultdict
from docx import Document
from merge import merge_group
from save import insert_blast_alignment_to_doc
from translate import translate_dna_to_protein



def main():
    # 按组名收集文件
    groups = defaultdict(list)
    for fname in os.listdir(input_dir):
        if fname.lower().endswith(("seq")):  # 这种写法更简洁
            # 使用 search 而不是 match，因为括号不一定在开头
            match = re.search(r"(\(.*?\))", fname)

            if match:
                group_name = match.group(1)
                groups[group_name].append(os.path.join(input_dir, fname))
                print(f"匹配成功: {group_name} -> {fname}")
            else:
                print(f"⚠️ 无法从文件名提取组名: {fname}")
    # Word 文档
    doc = Document()

    # 逐组处理
    for group, files in groups.items():
        final_seq = merge_group(files, group)
        if not final_seq:
            continue  # 缺文件的组跳过

        # 翻译成蛋白质
        protein_seq = translate_dna_to_protein(final_seq)

        # 输出组名和蛋白质序列
        doc.add_paragraph(group)
        doc.add_paragraph(protein_seq)
        doc.add_paragraph("")

        # 插入 BLAST 对比结果
        doc.add_paragraph("BLAST 2 Sequences 对比结果:")
        insert_blast_alignment_to_doc(doc, protein_seq, subject_seq)
        doc.add_paragraph("")  # 空行分隔

        print(f"✅ 已处理: {group}")

    # 保存 Word
    word_path = os.path.join(output_dir, output_doc_name)
    doc.save(word_path)
    print(f"📄 蛋白质序列及对比结果已保存到: {word_path}")


if __name__ == "__main__":
    # 运行前请执行: pip install python-docx requests beautifulsoup4
    # 输入目录
    input_dir = r"C:\Users\21017\Desktop\张翀_2270357268_测序结果"
    output_doc_name = "protein_sequences.docx"
    subject_seq = "MKPEDFRASTQRPFTGEEYLKSLQDGREIYIYGERVKDVTTHPAFRNAAASVAQLYDALHKPEMQDSLCWNTDTGSGGYTHKFFRVAKSADDLRQQRDAIAEWSRLSYGWMGRTPDYKAAFGCALGANPGFYGQFEQNARNWYTRIQETGLYFNHAIVNPPIDRHLPTDKVKDVYIKLEKETDAGIIVSGAKVVATNSALTHYNMIGFGSAQVMGENPDFALMFVAPMDADGVKLISRASYEMVAGATGSPYDYPLSSRFDENDAILVMDNVLIPWENVLIYRDFDRCRRWTMEGGFARMYPLQACVRLAVKLDFITALLKKSLECTGTLEFRGVQADLGEVVAWRNTFWALSDSMCSEATPWVNGAYLPDHAALQTYRVLAPMAYAKIKNIIERNVTSGLIYLPSSARDLNNPQIDQYLAKYVRGSNGMDHVQRIKILKLMWDAIGSEFGGRHELYEINYSGSQDEIRLQCLRQAQNSGNMDKMMAMVDRCLSEYDQDGWTVPHLHNNDDINMLDKLLK"
    output_dir = os.path.join(input_dir, "output")
    os.makedirs(output_dir, exist_ok=True)
    main()
