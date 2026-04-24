import time
from docx import Document
import requests
from bs4 import BeautifulSoup
from docx.shared import Pt, RGBColor

alignmentNotFoundError = ValueError("未找到 alignment 数据")


def upload(query: str, subject: str) -> str:
    """
    上传对比序列
    :param query: 查询序列
    :param subject: 对比序列
    :return: RID
    """
    url = "https://blast.ncbi.nlm.nih.gov/BlastAlign.cgi"

    data = {
        "QUERY": query,
        "SUBJECTS": subject,
        "db": "protein",
        "BL2SEQ": "on",
        "stype": "protein",
        "FORMAT_OBJECT": "Alignment",
        "FORMAT_TYPE": "HTML",
        "ALIGNMENT_VIEW": "Pairwise",
        "CMD": "request",
        "PROGRAM": "blastp"
    }

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                      "(KHTML, like Gecko) Chrome/140.0.0.0 Safari/537.36 Edg/140.0.0.0",
    }

    response = requests.post(url, data=data, headers=headers)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    rid_tag = soup.find("td", string="Request ID")  # 改用 string 参数
    if rid_tag:
        rid = rid_tag.find_next_sibling("td").b.text.strip()
        return rid
    else:
        raise ValueError("未找到 RID")


def download_alignment(rid: str) -> str:
    url = f"https://blast.ncbi.nlm.nih.gov/t2g.cgi?CMD=Get&RID={rid}"
    headers = {
        "accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "content-type": "application/x-www-form-urlencoded",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                      "(KHTML, like Gecko) Chrome/140.0.0.0 Safari/537.36 Edg/140.0.0.0",
        "referer": "https://blast.ncbi.nlm.nih.gov/BlastAlign.cgi"
    }

    resp = requests.get(url, headers=headers)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    div = soup.find("div", id="alignments", class_="blRes")
    if div and div.pre:
        return div.pre.get_text()
    else:
        raise alignmentNotFoundError



# 获取对比结果
def get_alignment(query_seq: str, subject_seq: str) -> str:
    rid = upload(query_seq, subject_seq)
    time.sleep(1)  # 这个接口是异步的以防万一还是要做一下等待
    for i in range(1, 3):
        try:
            alignment = download_alignment(rid)
            return alignment
        except alignmentNotFoundError:
            time.sleep(i * i)
        except Exception as e:
            raise ValueError(f"出现未知错误请检查程序/网络:{e}")
    raise ValueError("请检查网络或重试")



def insert_blast_alignment_to_doc(doc: Document, query_protein: str, subject_protein: str):
    """
  提交 BLAST 2 Sequences 请求并把对比结果插入 Word 文档表格
  并高亮突变位置，缩小字体，每组之间增加空行
  """
    try:
        pre_text = get_alignment(query_protein, subject_protein)
    except Exception as e:
        print(f"⚠️ {query_protein} 对比失败!:{e}")
        doc.add_paragraph("⚠️ BLAST 对比失败：" + str(e))
        return

    lines = pre_text.splitlines()

    # 提取统计信息
    stats_lines = []
    for line in lines:
        if line.strip().startswith("Score") or line.strip().startswith("Identities") or "Expect" in line:
            stats_lines.append(line.strip())
        elif line.strip().startswith("Query"):
            break

    if stats_lines:
        doc.add_paragraph("对比统计信息:")
        doc.add_paragraph(" | ".join(stats_lines))
        doc.add_paragraph("")  # 空行分隔

    # 添加表格标题
    doc.add_paragraph("BLAST 对比序列（表格显示）")

    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'

    i = 0
    not_first_group = False
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("Query"):
            query_parts = line.split()
            if len(query_parts) >= 4:
                if not_first_group:
                    # 插入空行
                    row = table.add_row().cells
                    row[0].text = ""
                    row[1].text = ""
                    row[2].text = ""
                    row[3].text = ""

                match_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
                sbjct_line = lines[i + 2].strip() if i + 2 < len(lines) else ""
                sbjct_parts = sbjct_line.split()

                # 插入 Query 行
                row = table.add_row().cells
                row[0].text = query_parts[0]
                row[1].text = query_parts[1]  # 起始位置
                row[2].text = query_parts[2]
                row[3].text = query_parts[3]

                # 插入 Match 行，高亮差异
                row = table.add_row().cells
                row[0].text = "Match"
                row[1].text = ""
                row[2].text = ""
                row[3].text = ""

                # 替换空格为数字空格
                replaced_line = match_line.replace(" ", "\u2007")
                p = row[2].paragraphs[0]
                run = p.add_run(replaced_line)
                run.font.size = Pt(8)

                # 如果有替换，就高亮整行
                if " " in match_line or "+" in match_line:
                    run.font.color.rgb = RGBColor(255, 0, 0)

                # 插入 Sbjct 行
                row = table.add_row().cells
                row[0].text = sbjct_parts[0]
                row[1].text = sbjct_parts[1]
                row[2].text = sbjct_parts[2]
                row[3].text = sbjct_parts[3]
                not_first_group = True

                i += 3
            else:
                i += 1
        else:
            i += 1

    # 缩小表格中字体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


